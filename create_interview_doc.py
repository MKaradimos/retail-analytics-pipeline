"""
Generate interview preparation document for the retail analytics pipeline project.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a heading with custom formatting."""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # EY-style dark blue
    return heading

def add_bullet_points(doc, items):
    """Add bullet points."""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)

def create_interview_prep_doc():
    """Create the interview preparation document."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Data Engineering Project', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Retail Analytics Pipeline - Interview Preparation Guide')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    # Add line break
    doc.add_paragraph()
    
    # ========================================
    # ELEVATOR PITCH
    # ========================================
    add_heading(doc, '🎯 Elevator Pitch (30 seconds)', 1)
    
    pitch = doc.add_paragraph()
    pitch_run = pitch.add_run(
        "I built a cloud-ready data engineering pipeline that consolidates data from multiple "
        "sources into a PostgreSQL data warehouse. The project demonstrates end-to-end ETL "
        "capabilities: API integration, data validation using Pydantic, transformation logic, "
        "and a star schema design for analytics. It's fully dockerized, includes automated data "
        "quality checks, and features pre-built analytics queries. The architecture is designed "
        "to scale and can easily be deployed to cloud platforms like AWS or Azure."
    )
    pitch_run.font.size = Pt(12)
    pitch_run.italic = True
    
    doc.add_paragraph()
    
    # ========================================
    # PROJECT OVERVIEW
    # ========================================
    add_heading(doc, '📊 Project Overview', 1)
    
    add_heading(doc, 'Business Context', 2)
    doc.add_paragraph(
        "A retail company needs to consolidate sales data from multiple sources, validate data "
        "quality, and enable business analytics for decision-making."
    )
    
    add_heading(doc, 'Technical Solution', 2)
    add_bullet_points(doc, [
        "Ingestion from 2 sources: REST API (products) and CSV files (sales transactions)",
        "Data validation layer using Pydantic models with business rules",
        "Transformation logic to convert raw data to warehouse-ready format",
        "PostgreSQL star schema (3 dimensions + 1 fact table)",
        "Docker containerization for portability",
        "Automated data quality checks and comprehensive logging"
    ])
    
    doc.add_page_break()
    
    # ========================================
    # TECHNICAL DEEP DIVE
    # ========================================
    add_heading(doc, '🔧 Technical Architecture', 1)
    
    add_heading(doc, 'Technology Stack', 2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    tech_stack = [
        ('Language', 'Python 3.11'),
        ('Data Validation', 'Pydantic'),
        ('Database', 'PostgreSQL 16'),
        ('API Client', 'Requests library'),
        ('Data Processing', 'Pandas'),
        ('Containerization', 'Docker & Docker Compose'),
        ('Database Admin', 'pgAdmin 4')
    ]
    
    for i, (component, tech) in enumerate(tech_stack):
        table.rows[i].cells[0].text = component
        table.rows[i].cells[1].text = tech
    
    doc.add_paragraph()
    
    add_heading(doc, 'Data Model - Star Schema', 2)
    doc.add_paragraph(
        "Implemented a dimensional modeling approach following Kimball methodology:"
    )
    
    add_bullet_points(doc, [
        "dim_product: Product master data (20 products from API)",
        "dim_customer: Customer master data (extracted from transactions)",
        "dim_date: Time dimension (pre-populated 2023-2025)",
        "fact_sales: Transaction-level sales data (500 sample transactions)"
    ])
    
    add_heading(doc, 'Key Design Decisions', 2)
    decisions_table = doc.add_table(rows=5, cols=2)
    decisions_table.style = 'Light List Accent 1'
    
    decisions = [
        ('Idempotency', 'Used ON CONFLICT clauses to handle reruns'),
        ('Error Handling', 'Retry logic with exponential backoff for API calls'),
        ('Validation', 'Pydantic models enforce data types and business rules'),
        ('Logging', 'Structured logging with timestamps and severity levels'),
        ('Modularity', 'Separated ingestion, transformation, and loading layers')
    ]
    
    for i, (decision, rationale) in enumerate(decisions):
        decisions_table.rows[i].cells[0].text = decision
        decisions_table.rows[i].cells[1].text = rationale
    
    doc.add_page_break()
    
    # ========================================
    # INTERVIEW QUESTIONS
    # ========================================
    add_heading(doc, '💬 Expected Interview Questions & Answers', 1)
    
    questions = [
        {
            'q': 'Why did you choose a star schema?',
            'a': "Star schemas are optimized for analytical queries. They provide fast query "
                 "performance through denormalization and are easy for business users to understand. "
                 "The dimensional model naturally supports slicing and dicing data by time, product, "
                 "and customer - which are the key analytics dimensions for retail."
        },
        {
            'q': 'How do you handle data quality issues?',
            'a': "I implemented multiple layers: (1) Pydantic models for schema validation and type "
                 "checking, (2) Business rule validation (e.g., prices must be positive), "
                 "(3) Automated quality checks after loading (orphan detection, negative amounts, "
                 "date consistency), (4) Comprehensive logging to track validation failures."
        },
        {
            'q': 'What would you do differently in production?',
            'a': "Several enhancements: (1) Use Apache Airflow for workflow orchestration, "
                 "(2) Implement incremental loading instead of full refreshes, (3) Add monitoring "
                 "and alerting (e.g., Datadog), (4) Implement proper secrets management (e.g., AWS "
                 "Secrets Manager), (5) Add comprehensive unit and integration tests, (6) Implement "
                 "data lineage tracking, (7) Add table partitioning for large fact tables."
        },
        {
            'q': 'How would you scale this to handle millions of records?',
            'a': "Key strategies: (1) Implement batch processing with configurable batch sizes, "
                 "(2) Use connection pooling for database efficiency, (3) Add table partitioning "
                 "by date, (4) Implement parallel processing for independent tasks, (5) Consider "
                 "cloud data warehouses like Snowflake or BigQuery for large-scale analytics, "
                 "(6) Use columnar storage formats like Parquet for data lakes."
        },
        {
            'q': 'Why Docker?',
            'a': "Docker ensures consistency across environments - development, testing, and "
                 "production. It eliminates 'works on my machine' problems and makes deployment "
                 "straightforward. The docker-compose setup orchestrates multiple services "
                 "(PostgreSQL, pipeline, pgAdmin) making it easy to spin up the entire stack. "
                 "It's also cloud-agnostic and works with ECS, Kubernetes, etc."
        },
        {
            'q': 'How do you ensure data consistency?',
            'a': "Multiple mechanisms: (1) Database transactions for atomicity, (2) Foreign key "
                 "constraints to maintain referential integrity, (3) Idempotent operations using "
                 "ON CONFLICT, (4) Automated validation checks after loading, (5) Logging all "
                 "operations for audit trail."
        },
        {
            'q': 'What analytics can you perform with this warehouse?',
            'a': "I've implemented 14 pre-built analytics queries including: sales performance by "
                 "product and category, customer segmentation and lifetime value, sales trends "
                 "(daily, weekly, monthly), location-based analysis, payment method analysis, "
                 "weekend vs weekday comparisons, and growth rate calculations. The star schema "
                 "makes it easy to create custom queries for ad-hoc analysis."
        }
    ]
    
    for i, qa in enumerate(questions, 1):
        add_heading(doc, f"Q{i}: {qa['q']}", 3)
        answer = doc.add_paragraph(qa['a'])
        answer.paragraph_format.space_after = Pt(12)
    
    doc.add_page_break()
    
    # ========================================
    # TECHNICAL SPECIFICS
    # ========================================
    add_heading(doc, '📝 Technical Implementation Details', 1)
    
    add_heading(doc, 'Pipeline Flow', 2)
    add_bullet_points(doc, [
        "Step 1: Initialize database schema (creates tables, indexes, views)",
        "Step 2: Fetch products from FakeStore API with retry logic",
        "Step 3: Validate products using Pydantic models",
        "Step 4: Transform to dimension table format",
        "Step 5: Load to dim_product table (upsert logic)",
        "Step 6: Read sales transactions from CSV",
        "Step 7: Validate transactions and extract customer dimension",
        "Step 8: Load customers and transactions to warehouse",
        "Step 9: Run automated data quality checks",
        "Step 10: Generate execution summary with statistics"
    ])
    
    add_heading(doc, 'Code Highlights', 2)
    
    doc.add_paragraph("Context Manager for Database Connections:", style='Intense Quote')
    code1 = doc.add_paragraph(
        "@contextmanager\n"
        "def get_connection():\n"
        "    conn = psycopg2.connect(**DB_CONFIG)\n"
        "    try:\n"
        "        yield conn\n"
        "        conn.commit()\n"
        "    except:\n"
        "        conn.rollback()\n"
        "        raise\n"
        "    finally:\n"
        "        conn.close()"
    )
    code1.runs[0].font.name = 'Courier New'
    code1.runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    
    doc.add_paragraph("Pydantic Validation Example:", style='Intense Quote')
    code2 = doc.add_paragraph(
        "class SalesTransactionModel(BaseModel):\n"
        "    transaction_id: str\n"
        "    quantity: int\n"
        "    total_amount: Decimal\n"
        "    \n"
        "    @validator('quantity')\n"
        "    def quantity_positive(cls, v):\n"
        "        if v <= 0:\n"
        "            raise ValueError('Quantity must be positive')\n"
        "        return v"
    )
    code2.runs[0].font.name = 'Courier New'
    code2.runs[0].font.size = Pt(9)
    
    doc.add_page_break()
    
    # ========================================
    # DEMONSTRATED SKILLS
    # ========================================
    add_heading(doc, '✅ Skills Demonstrated', 1)
    
    skills_table = doc.add_table(rows=10, cols=2)
    skills_table.style = 'Light Grid Accent 1'
    
    skills = [
        ('Data Engineering', 'ETL pipeline design and implementation'),
        ('Data Modeling', 'Star schema / dimensional modeling'),
        ('Data Quality', 'Validation, quality checks, error handling'),
        ('API Integration', 'REST API consumption with retry logic'),
        ('Python', 'OOP, context managers, type hints, Pydantic'),
        ('SQL', 'DDL, DML, views, indexes, constraints'),
        ('Docker', 'Containerization, docker-compose, multi-service setup'),
        ('DevOps', 'Configuration management, logging, Makefile'),
        ('Documentation', 'Comprehensive README, code comments, docstrings'),
        ('Best Practices', 'Modularity, error handling, idempotency, testing mindset')
    ]
    
    for i, (skill, detail) in enumerate(skills):
        skills_table.rows[i].cells[0].text = skill
        skills_table.rows[i].cells[1].text = detail
    
    doc.add_paragraph()
    
    # ========================================
    # DEMO WALKTHROUGH
    # ========================================
    add_heading(doc, '🖥️ Demo Walkthrough (5 minutes)', 1)
    
    demo_steps = [
        "1. Architecture Overview (30 sec): Show README architecture diagram",
        "2. Code Structure (1 min): Quick tour of ingestion/, transformations/, sql/",
        "3. Data Validation (1 min): Show Pydantic models and validation logic",
        "4. Docker Setup (1 min): docker-compose.yml and Makefile",
        "5. Pipeline Execution (1 min): Run 'make run' and show logs",
        "6. Database Queries (1.5 min): Connect to PostgreSQL and run analytics queries"
    ]
    
    for step in demo_steps:
        p = doc.add_paragraph(step)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    
    add_heading(doc, 'Sample Analytics Query to Show', 2)
    query = doc.add_paragraph(
        "-- Top 5 products by revenue\n"
        "SELECT product_name, category, total_revenue\n"
        "FROM vw_product_sales_summary\n"
        "ORDER BY total_revenue DESC\n"
        "LIMIT 5;"
    )
    query.runs[0].font.name = 'Courier New'
    query.runs[0].font.size = Pt(10)
    
    doc.add_page_break()
    
    # ========================================
    # EY ALIGNMENT
    # ========================================
    add_heading(doc, '🎯 EY Alignment', 1)
    
    doc.add_paragraph(
        "This project directly aligns with EY's data engineering and analytics consulting work:"
    )
    
    add_bullet_points(doc, [
        "Client Scenario: Mirrors real client engagements for data consolidation",
        "Cloud-Ready: Architecture designed for AWS/Azure/GCP deployment",
        "Consulting Mindset: Documentation explains business value, not just tech",
        "Best Practices: Follows industry standards (star schema, data validation, logging)",
        "Scalability: Discusses production enhancements and scaling strategies",
        "Tool Proficiency: Uses tools common in EY projects (Python, PostgreSQL, Docker)"
    ])
    
    add_heading(doc, 'Talking Points for EY', 2)
    add_bullet_points(doc, [
        "\"I designed this with consulting delivery in mind - clear documentation, "
        "reproducible setup, and business-focused analytics.\"",
        
        "\"The modular architecture makes it easy to add new data sources as client "
        "requirements evolve.\"",
        
        "\"I included automated quality checks because data trust is critical in "
        "consulting engagements.\"",
        
        "\"The Docker setup ensures consistent deployments across different client "
        "environments.\"",
        
        "\"I documented future enhancements to show I think about production readiness "
        "and long-term sustainability.\""
    ])
    
    doc.add_paragraph()
    
    # ========================================
    # QUICK REFERENCE
    # ========================================
    add_heading(doc, '📋 Quick Reference', 1)
    
    add_heading(doc, 'GitHub Repository', 2)
    doc.add_paragraph('[Your GitHub URL]', style='Intense Quote')
    
    add_heading(doc, 'Key Files to Mention', 2)
    add_bullet_points(doc, [
        'README.md - Comprehensive project documentation',
        'pipeline.py - Main ETL orchestrator',
        'models.py - Pydantic validation models',
        'sql/schema.sql - Data warehouse DDL',
        'sql/analytics_queries.sql - 14 pre-built analytics queries',
        'docker-compose.yml - Multi-service setup'
    ])
    
    add_heading(doc, 'Key Metrics', 2)
    metrics_table = doc.add_table(rows=4, cols=2)
    metrics_table.style = 'Medium Grid 1 Accent 1'
    
    metrics = [
        ('Products in Warehouse', '20 (from API)'),
        ('Sample Transactions', '500 (generated)'),
        ('Data Quality Checks', '3 automated checks'),
        ('Analytics Queries', '14 pre-built queries')
    ]
    
    for i, (metric, value) in enumerate(metrics):
        metrics_table.rows[i].cells[0].text = metric
        metrics_table.rows[i].cells[1].text = value
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f'\nDocument Generated: {datetime.now().strftime("%B %d, %Y")}')
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    output_path = '/home/claude/retail-analytics-pipeline/docs/interview_preparation.docx'
    doc.save(output_path)
    print(f"✓ Interview preparation document created: {output_path}")
    
    return output_path

if __name__ == "__main__":
    create_interview_prep_doc()
